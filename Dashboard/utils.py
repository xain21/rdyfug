import pytesseract
from PIL import Image, ImageEnhance
import fitz  # PyMuPDF
import concurrent.futures

# Preprocess image (grayscale + contrast enhancement)
def preprocess_image(image):
    grayscale_img = image.convert("L")  # Convert to grayscale
    enhancer = ImageEnhance.Contrast(grayscale_img)
    enhanced_img = enhancer.enhance(2)  # Increase contrast
    return enhanced_img

# Process each page and optionally apply OCR (use_ocr=False skips Tesseract
# entirely for a much faster "text only" search).
def process_page_ocr(page, pdf_path, meeting_type, use_ocr=True):
    # Get normal PDF text
    page_text = page.get_text()

    if not use_ocr:
        return page_text

    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

    # OCR the page too (used when the caller explicitly wants OCR + text)
    try:
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        preprocessed_img = preprocess_image(img)

        ocr_text = pytesseract.image_to_string(preprocessed_img)

        # Combine normal PDF text + OCR text
        return page_text + "\n" + ocr_text

    except Exception as e:
        print(f"OCR error on page {page.number + 1}: {e}")
        return page_text

# Extract text from a PDF using OCR in parallel (for BOD meetings only)
def extract_text_with_ocr_parallel(pdf_path, meeting_type):
    text = ""
    with fitz.open(pdf_path) as pdf:
        with concurrent.futures.ThreadPoolExecutor() as executor:
            futures = []
            for page in pdf:
                # Submit each page for OCR processing only for BOD meetings
                futures.append(executor.submit(process_page_ocr, page, pdf_path, meeting_type))

            # Collect results as they complete
            for future in concurrent.futures.as_completed(futures):
                text += future.result()  # Add the result from OCR or text extraction
    return text



import fitz  # PyMuPDF

def highlight_pdf(pdf_path, search_term):
    """
    Highlight the search term in the provided PDF and return the path of the modified PDF.
    """
    with fitz.open(pdf_path) as doc:
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)  # Load the page
            text_instances = page.search_for(search_term)  # Search for the term

            # Highlight the found text instances
            for inst in text_instances:
                page.add_highlight_annot(inst)  # Highlight the text instance

        # Save the modified PDF with highlights
        highlighted_pdf_path = pdf_path.replace("meeting_pdfs", "highlighted_pdfs")
        doc.save(highlighted_pdf_path)
        return highlighted_pdf_path


# ============================================================================
# AUTO-EXTRACTION: read an uploaded PDF/Word file and pull out a meeting
# title, date, and agenda items automatically, instead of typing them by hand.
# ============================================================================

import os
import re
import json


def extract_text_from_file(file_path):
    """
    Extract raw text from a PDF or Word (.docx) file.
    - Text-based PDFs: read directly with PyMuPDF (fast).
    - Scanned/image PDFs: automatically falls back to OCR (pytesseract).
    - .docx: read with python-docx.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.docx':
        try:
            import docx  # python-docx
        except ImportError as e:
            raise RuntimeError(
                "python-docx is not installed. Run: pip install python-docx"
            ) from e
        document = docx.Document(file_path)
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        # Also pull text out of any tables (agendas are sometimes tabular)
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    elif ext == '.pdf':
        text = ""
        with fitz.open(file_path) as pdf:
            for page in pdf:
                text += page.get_text() + "\n"

        # Almost no real text extracted -> this is very likely a scanned
        # (image-only) PDF, so fall back to OCR automatically.
        if len(text.strip()) < 40:
            text = extract_text_with_ocr_parallel(file_path, meeting_type=None)

        return text

    else:
        raise ValueError(f"Unsupported file type: {ext}. Only .pdf and .docx are supported.")


_DATE_REGEX = re.compile(
    r'(\d{1,2}(?:st|nd|rd|th)?\s+[A-Za-z]{3,9}\s+\d{4}'   # 18th August 2026
    r'|[A-Za-z]{3,9}\s+\d{1,2},?\s+\d{4}'                  # August 18, 2026
    r'|\d{1,2}[/-]\d{1,2}[/-]\d{2,4})'                      # 18/08/2026
)


def _find_date(text):
    """Best-effort date guess: looks for 'held on <date>' first, then any
    date-shaped string near the top of the document."""
    try:
        from dateutil import parser as dateparser
    except ImportError:
        return None

    candidates = []
    m = re.search(r'held on\s+(.+?)(?:[.\n]|$)', text, re.IGNORECASE)
    if m:
        candidates.append(m.group(1))
    candidates += _DATE_REGEX.findall(text[:800])

    for c in candidates:
        try:
            return dateparser.parse(c, fuzzy=True, dayfirst=False).date()
        except Exception:
            continue
    return None


def parse_meeting_document_regex(text):
    """
    Rule-based extraction. Works well when the document follows a
    predictable layout (a numbered "Agenda" list, like your finance
    committee MOM template). This is the free, no-API fallback.
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    title = lines[0] if lines else "Untitled Meeting"

    date = _find_date(text)

    agenda_items = []
    m = re.search(
        r'Agenda\s*\n(.*?)(?:\n\s*(?:Discussion|Decision|Minutes|Action\s*Items?)\b|$)',
        text, re.IGNORECASE | re.DOTALL,
    )
    if m:
        block = m.group(1)
        items = re.findall(r'^\s*\d+[\.\)]\s*(.+)$', block, re.MULTILINE)
        agenda_items = [i.strip() for i in items if i.strip()]

    discussion = ""
    m2 = re.search(
        r'Discussion\s*\n(.*?)(?:\n\s*Decision[s]?\b|$)',
        text, re.IGNORECASE | re.DOTALL,
    )
    if m2:
        discussion = m2.group(1).strip()

    decision = ""
    m3 = re.search(r'Decision[s]?\s*\n(.*)', text, re.IGNORECASE | re.DOTALL)
    if m3:
        decision = m3.group(1).strip()

    return {
        "title": title,
        "date": date.isoformat() if date else "",
        "agenda_items": agenda_items,
        "discussion": discussion,
        "decision": decision,
        "source": "regex",
    }


def parse_meeting_document_ai(text, api_key=None):
    """
    AI-based extraction (Anthropic Claude API). Handles messy formatting,
    OCR noise, or documents that don't follow one fixed template - things
    the regex parser above can't. Returns None (silently) if no API key is
    configured or the call fails, so callers can fall back to regex.
    """
    api_key = api_key or os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return None

    try:
        import requests
    except ImportError:
        return None

    prompt = (
        "You are extracting structured data from a board/committee meeting "
        "document. Read the text below and return ONLY a raw JSON object "
        "(no markdown fences, no commentary) with these exact keys:\n"
        '  "title": short meeting title (string)\n'
        '  "date": meeting date in YYYY-MM-DD format if found, else ""\n'
        '  "agenda_items": array of agenda item strings, in order\n'
        '  "discussion": a short plain-text summary of the discussion (string)\n'
        '  "decision": a short plain-text summary of decisions/action items (string)\n\n'
        "Document text:\n---\n" + text[:12000] + "\n---"
    )

    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            json={
                "model": "claude-sonnet-4-6",
                "max_tokens": 2000,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=60,
        )
        resp.raise_for_status()
        data = resp.json()
        raw = "".join(
            block.get("text", "")
            for block in data.get("content", [])
            if block.get("type") == "text"
        ).strip()

        if raw.startswith("```"):
            raw = re.sub(r'^```(?:json)?|```$', '', raw, flags=re.MULTILINE).strip()

        parsed = json.loads(raw)
        parsed.setdefault("title", "Untitled Meeting")
        parsed.setdefault("date", "")
        parsed.setdefault("agenda_items", [])
        parsed.setdefault("discussion", "")
        parsed.setdefault("decision", "")
        parsed["source"] = "ai"
        return parsed
    except Exception as e:
        print(f"AI meeting extraction failed, falling back to regex: {e}")
        return None


def parse_meeting_document(text):
    """
    Hybrid entry point used by the upload view: try AI extraction first
    (only runs if ANTHROPIC_API_KEY is set in .env), fall back to the free
    regex parser if no key is configured or the AI call fails/returns junk.
    """
    ai_result = parse_meeting_document_ai(text)
    if ai_result and ai_result.get("title") and ai_result.get("agenda_items"):
        return ai_result
    return parse_meeting_document_regex(text)
